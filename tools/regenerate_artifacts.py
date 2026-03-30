#!/usr/bin/env python3
"""Regenerate document and spreadsheet artifacts from markdown and CSV sources."""

from __future__ import annotations

import csv
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Iterable, List, Sequence
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Image as RLImage
from reportlab.platypus import KeepTogether, LongTable, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle

from analyze_expanded_dataset import main as analyze_expanded_dataset
from generate_figure_set import main as generate_figure_set


ROOT = Path(__file__).resolve().parent.parent
DOCUMENT_AUTHOR = "Raj Anand Raghavan"

MARKDOWN_ARTIFACTS = [
    {
        "source": ROOT / "Research_Paper_FINAL.md",
        "docx": ROOT / "Research_Paper_FINAL.docx",
        "pdf": ROOT / "Research_Paper_FINAL.pdf",
    },
    {
        "source": ROOT / "Research_Paper_FINAL.md",
        "docx": ROOT / "Research_Paper_COMPLETE_INTEGRATED.docx",
        "pdf": None,
    },
    {
        "source": ROOT / "Research_Paper_MDPI_Strengthened.md",
        "docx": ROOT / "Research_Paper_MDPI_Strengthened.docx",
        "pdf": ROOT / "Research_Paper_MDPI_Strengthened.pdf",
    },
    {
        "source": ROOT / "Supplementary_Materials_Statistical_Strengthening.md",
        "docx": ROOT / "Supplementary_Materials_Statistical_Strengthening.docx",
        "pdf": ROOT / "Supplementary_Materials_Statistical_Strengthening.pdf",
    },
]

CSV_XLSX_ARTIFACTS = [
    (ROOT / "analysis_data.csv", ROOT / "analysis_data.xlsx"),
    (ROOT / "statistical_analysis_results.csv", ROOT / "statistical_analysis_results.xlsx"),
    (ROOT / "monte_carlo_results.csv", ROOT / "monte_carlo_results.xlsx"),
    (ROOT / "enhanced_analysis_results.csv", ROOT / "enhanced_analysis_results.xlsx"),
    (ROOT / "expanded_dataset_40_models.csv", ROOT / "expanded_dataset_40_models.xlsx"),
]


ESCAPED_MARKDOWN_RE = re.compile(r"\\([\\`*_{}\[\]()#+\-.!|])")
INLINE_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*\n]+?\*)")
UNORDERED_LIST_RE = re.compile(r"^\s*-\s+(.*)$")
ORDERED_LIST_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?[\s:\-|]+\|?\s*$")
IMAGE_BLOCK_RE = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")
SUBSCRIPT_TRANSLATION = str.maketrans("₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎", "0123456789+-=()")


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_default_docx_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True

    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)


def normalize_inline_text(text: str) -> str:
    text = ESCAPED_MARKDOWN_RE.sub(r"\1", text)
    return text.translate(SUBSCRIPT_TRANSLATION)


def document_title(source: Path) -> str:
    blocks = parse_markdown_blocks(source.read_text(encoding="utf-8"))
    for block in blocks:
        if block["type"] == "heading" and block["level"] == 1:
            return normalize_inline_text(block["text"])
    return normalize_inline_text(source.stem.replace("_", " "))


def document_subject(source: Path) -> str:
    if source.name == "Research_Paper_MDPI_Strengthened.md":
        return "Athletic footwear screening-LCA manuscript"
    if source.name == "Supplementary_Materials_Statistical_Strengthening.md":
        return "Supplementary materials for athletic footwear screening-LCA manuscript"
    return "Research manuscript artifact"


def parse_inline_chunks(text: str) -> List[dict]:
    text = normalize_inline_text(text)
    chunks: List[dict] = []
    cursor = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            chunks.append({"text": text[cursor:match.start()], "bold": False, "italic": False, "code": False})
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            chunks.append({"text": token[2:-2], "bold": True, "italic": False, "code": False})
        elif token.startswith("`") and token.endswith("`"):
            chunks.append({"text": token[1:-1], "bold": False, "italic": False, "code": True})
        else:
            chunks.append({"text": token[1:-1], "bold": False, "italic": True, "code": False})
        cursor = match.end()
    if cursor < len(text):
        chunks.append({"text": text[cursor:], "bold": False, "italic": False, "code": False})
    return [chunk for chunk in chunks if chunk["text"]]


def render_docx_runs(paragraph, text: str) -> None:
    for chunk in parse_inline_chunks(text):
        run = paragraph.add_run(chunk["text"])
        run.bold = chunk["bold"]
        run.italic = chunk["italic"]
        if chunk["code"]:
            run.font.name = "Courier New"


def to_reportlab_markup(text: str) -> str:
    parts = []
    for chunk in parse_inline_chunks(text):
        escaped = escape(chunk["text"])
        if chunk["code"]:
            parts.append(f'<font face="Courier">{escaped}</font>')
        elif chunk["bold"]:
            parts.append(f"<b>{escaped}</b>")
        elif chunk["italic"]:
            parts.append(f"<i>{escaped}</i>")
        else:
            parts.append(escaped)
    return "".join(parts)


def to_html_markup(text: str) -> str:
    parts = []
    for chunk in parse_inline_chunks(text):
        escaped = escape(chunk["text"])
        if chunk["code"]:
            parts.append(f"<code>{escaped}</code>")
        elif chunk["bold"]:
            parts.append(f"<strong>{escaped}</strong>")
        elif chunk["italic"]:
            parts.append(f"<em>{escaped}</em>")
        else:
            parts.append(escaped)
    return "".join(parts)


def is_table_start(lines: Sequence[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    separator = lines[idx + 1].strip()
    return (
        lines[idx].strip().startswith("|")
        and "-" in separator
        and TABLE_SEPARATOR_RE.match(separator) is not None
    )


def split_table_row(line: str) -> List[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def parse_image_block(line: str) -> dict | None:
    match = IMAGE_BLOCK_RE.match(line)
    if not match:
        return None
    path = match.group(2).strip()
    if path.startswith("<") and path.endswith(">"):
        path = path[1:-1]
    return {"type": "image", "alt": normalize_inline_text(match.group(1).strip()), "path": path}


def resolve_asset_path(source: Path, asset: str) -> Path:
    return (source.parent / asset).resolve()


def parse_markdown_blocks(text: str) -> List[dict]:
    lines = text.splitlines()
    blocks: List[dict] = []
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("```"):
            idx += 1
            code_lines = []
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            idx += 1
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            continue

        if stripped == "---":
            blocks.append({"type": "rule"})
            idx += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text_value = stripped[level:].strip()
            blocks.append({"type": "heading", "level": min(level, 3), "text": text_value})
            idx += 1
            continue

        image_block = parse_image_block(stripped)
        if image_block is not None:
            blocks.append(image_block)
            idx += 1
            continue

        if is_table_start(lines, idx):
            header = split_table_row(lines[idx])
            idx += 2
            rows = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                rows.append(split_table_row(lines[idx]))
                idx += 1
            blocks.append({"type": "table", "header": header, "rows": rows})
            continue

        unordered = UNORDERED_LIST_RE.match(line)
        ordered = ORDERED_LIST_RE.match(line)
        if unordered or ordered:
            ordered_list = ordered is not None
            items = []
            while idx < len(lines):
                current = lines[idx]
                match = ORDERED_LIST_RE.match(current) if ordered_list else UNORDERED_LIST_RE.match(current)
                if not match:
                    break
                item = match.group(1).strip()
                idx += 1
                continuation = []
                while idx < len(lines):
                    nxt = lines[idx]
                    if (
                        not nxt.strip()
                        or nxt.strip().startswith("#")
                        or nxt.strip() == "---"
                        or nxt.strip().startswith("|")
                        or UNORDERED_LIST_RE.match(nxt)
                        or ORDERED_LIST_RE.match(nxt)
                        or nxt.strip().startswith("```")
                    ):
                        break
                    continuation.append(nxt.strip())
                    idx += 1
                if continuation:
                    item = " ".join([item] + continuation)
                items.append(item)
            blocks.append({"type": "list", "ordered": ordered_list, "items": items})
            continue

        paragraph_lines = [stripped]
        idx += 1
        while idx < len(lines):
            nxt = lines[idx]
            nxt_stripped = nxt.strip()
            if (
                not nxt_stripped
                or nxt_stripped.startswith("#")
                or nxt_stripped == "---"
                or nxt_stripped.startswith("```")
                or is_table_start(lines, idx)
                or UNORDERED_LIST_RE.match(nxt)
                or ORDERED_LIST_RE.match(nxt)
            ):
                break
            paragraph_lines.append(nxt_stripped)
            idx += 1
        blocks.append({"type": "paragraph", "text": " ".join(paragraph_lines)})
    return blocks


def markdown_blocks_to_html(blocks: Sequence[dict], title: str, source: Path) -> str:
    body_parts: List[str] = []
    for block in blocks:
        btype = block["type"]
        if btype == "heading":
            level = max(1, min(block["level"], 3))
            body_parts.append(f"<h{level}>{to_html_markup(block['text'])}</h{level}>")
        elif btype == "paragraph":
            body_parts.append(f"<p>{to_html_markup(block['text'])}</p>")
        elif btype == "list":
            tag = "ol" if block["ordered"] else "ul"
            items = "".join(f"<li>{to_html_markup(item)}</li>" for item in block["items"])
            body_parts.append(f"<{tag}>{items}</{tag}>")
        elif btype == "code":
            body_parts.append(f"<pre><code>{escape(normalize_inline_text(block['text']))}</code></pre>")
        elif btype == "table":
            header_cells = "".join(f"<th>{to_html_markup(cell)}</th>" for cell in block["header"])
            body_rows = []
            for row in block["rows"]:
                cells = "".join(f"<td>{to_html_markup(cell)}</td>" for cell in row)
                body_rows.append(f"<tr>{cells}</tr>")
            body_parts.append(
                "<table><thead><tr>"
                + header_cells
                + "</tr></thead><tbody>"
                + "".join(body_rows)
                + "</tbody></table>"
            )
        elif btype == "image":
            image_path = resolve_asset_path(source, block["path"])
            if image_path.exists():
                body_parts.append(
                    "<figure>"
                    f"<img src=\"{escape(image_path.as_uri())}\" alt=\"{escape(block['alt'])}\">"
                    f"<figcaption>{to_html_markup(block['alt'])}</figcaption>"
                    "</figure>"
                )
            else:
                body_parts.append(f"<p><em>Missing figure asset:</em> {escape(block['path'])}</p>")
        elif btype == "rule":
            body_parts.append("<hr>")

    stylesheet = """
    @page { size: Letter; margin: 0.75in; }
    html { -webkit-print-color-adjust: exact; print-color-adjust: exact; }
    body { font-family: "Times New Roman", Times, serif; font-size: 10.5pt; line-height: 1.35; color: #111; }
    h1 { font-size: 18pt; line-height: 1.2; text-align: center; margin: 0 0 10pt; page-break-after: avoid; }
    h2 { font-size: 13pt; line-height: 1.25; margin: 12pt 0 6pt; page-break-after: avoid; }
    h3 { font-size: 11.5pt; line-height: 1.25; margin: 10pt 0 5pt; page-break-after: avoid; }
    p { margin: 0 0 8pt; orphans: 3; widows: 3; }
    ul, ol { margin: 0 0 8pt 1.15rem; padding: 0; }
    li { margin: 0 0 4pt; }
    pre { font-family: "Courier New", Courier, monospace; font-size: 8.5pt; line-height: 1.3; white-space: pre-wrap; border: 1px solid #bbb; padding: 8pt; margin: 0 0 8pt; }
    code { font-family: "Courier New", Courier, monospace; font-size: 0.95em; }
    hr { border: 0; border-top: 1px solid #bbb; margin: 8pt 0; }
    table { width: 100%; border-collapse: collapse; table-layout: fixed; font-size: 8.4pt; margin: 0 0 10pt; break-inside: auto; page-break-inside: auto; }
    thead { display: table-header-group; }
    tbody { display: table-row-group; }
    tr { break-inside: avoid; page-break-inside: avoid; }
    th, td { border: 1px solid #999; padding: 4pt; text-align: left; vertical-align: top; overflow-wrap: anywhere; word-break: break-word; }
    th { background: #d9eaf7; font-weight: bold; }
    tbody tr:nth-child(even) { background: #f8fbfd; }
    strong { font-weight: 700; }
    em { font-style: italic; }
    figure { margin: 10pt 0 14pt; break-inside: avoid; page-break-inside: avoid; }
    figure img { display: block; max-width: 100%; max-height: 8.2in; width: auto; height: auto; margin: 0 auto 6pt; object-fit: contain; }
    figcaption { font-size: 9pt; line-height: 1.25; text-align: center; font-style: italic; }
    """
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<title>{escape(title)}</title>"
        f"<meta name='author' content='{escape(DOCUMENT_AUTHOR)}'>"
        f"<meta name='description' content='{escape(document_subject(source))}'>"
        f"<style>{stylesheet}</style></head><body>"
        + "".join(body_parts)
        + "</body></html>"
    )


def find_chrome_binary() -> str | None:
    candidates = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        shutil.which("google-chrome"),
        shutil.which("chrome"),
        shutil.which("chromium"),
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    return None


def render_markdown_to_docx(source: Path, output: Path) -> None:
    document = Document()
    set_default_docx_style(document)
    title = document_title(source)
    document.core_properties.title = title
    document.core_properties.author = DOCUMENT_AUTHOR
    document.core_properties.subject = document_subject(source)
    document.core_properties.comments = "Generated from markdown source"
    document.core_properties.last_modified_by = DOCUMENT_AUTHOR
    document.core_properties.keywords = "footwear, LCA, sustainability, manuscript"

    blocks = parse_markdown_blocks(source.read_text(encoding="utf-8"))

    for block in blocks:
        btype = block["type"]
        if btype == "heading":
            paragraph = document.add_paragraph(style=f"Heading {block['level']}")
            render_docx_runs(paragraph, block["text"])
        elif btype == "paragraph":
            paragraph = document.add_paragraph()
            render_docx_runs(paragraph, block["text"])
            paragraph.paragraph_format.space_after = Pt(8)
        elif btype == "list":
            style_name = "List Number" if block["ordered"] else "List Bullet"
            for item in block["items"]:
                paragraph = document.add_paragraph(style=style_name)
                render_docx_runs(paragraph, item)
        elif btype == "code":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(block["text"])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif btype == "table":
            rows = [block["header"]] + block["rows"]
            table = document.add_table(rows=len(rows), cols=len(block["header"]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for row_idx, row in enumerate(rows):
                for col_idx, value in enumerate(row):
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    render_docx_runs(paragraph, value)
                    if row_idx == 0:
                        set_cell_shading(cell, "D9EAF7")
                        for run in paragraph.runs:
                            run.bold = True
            document.add_paragraph()
        elif btype == "image":
            image_path = resolve_asset_path(source, block["path"])
            if image_path.exists():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.add_run().add_picture(str(image_path), width=Inches(6.4))
                caption = document.add_paragraph()
                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption.paragraph_format.space_after = Pt(10)
                render_docx_runs(caption, block["alt"])
                for run in caption.runs:
                    run.italic = True
            else:
                paragraph = document.add_paragraph()
                render_docx_runs(paragraph, f"Missing figure asset: {block['path']}")
        elif btype == "rule":
            document.add_paragraph()

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def build_pdf_styles():
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=10.5,
        leading=14,
        spaceAfter=8,
    )
    heading1 = ParagraphStyle(
        "Heading1Custom",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=18,
        leading=22,
        spaceAfter=10,
        alignment=TA_CENTER,
    )
    heading2 = ParagraphStyle(
        "Heading2Custom",
        parent=styles["Heading2"],
        fontName="Times-Bold",
        fontSize=13,
        leading=16,
        spaceAfter=8,
        spaceBefore=4,
    )
    heading3 = ParagraphStyle(
        "Heading3Custom",
        parent=styles["Heading3"],
        fontName="Times-Bold",
        fontSize=11.5,
        leading=14,
        spaceAfter=6,
        spaceBefore=2,
    )
    bullet = ParagraphStyle(
        "BulletCustom",
        parent=body,
        leftIndent=18,
        firstLineIndent=-10,
        spaceAfter=4,
    )
    code = ParagraphStyle(
        "CodeCustom",
        parent=body,
        fontName="Courier",
        fontSize=8.5,
        leading=10,
        leftIndent=12,
        rightIndent=12,
    )
    caption = ParagraphStyle(
        "CaptionCustom",
        parent=body,
        fontName="Times-Italic",
        fontSize=9,
        leading=11,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    table_header = ParagraphStyle(
        "TableHeader",
        parent=body,
        fontName="Times-Bold",
        fontSize=8.2,
        leading=9.5,
        spaceAfter=0,
        wordWrap="CJK",
    )
    table_cell = ParagraphStyle(
        "TableCell",
        parent=body,
        fontName="Times-Roman",
        fontSize=8.2,
        leading=9.4,
        spaceAfter=0,
        wordWrap="CJK",
    )
    return {
        "body": body,
        "heading1": heading1,
        "heading2": heading2,
        "heading3": heading3,
        "bullet": bullet,
        "code": code,
        "caption": caption,
        "table_header": table_header,
        "table_cell": table_cell,
    }


def approximate_table_widths(rows: Sequence[Sequence[str]], available_width: float) -> List[float]:
    col_count = max(len(row) for row in rows)
    if col_count == 1:
        return [available_width]
    lengths = [0] * col_count
    averages = [0.0] * col_count
    for row in rows:
        for idx, cell in enumerate(row):
            text_length = len(normalize_inline_text(cell))
            lengths[idx] = max(lengths[idx], text_length)
            averages[idx] += text_length
    count = max(len(rows), 1)
    weights = []
    for idx, length in enumerate(lengths):
        average = averages[idx] / count
        weights.append(max(min((length * 0.6) + (average * 0.4), 45), 8))
    if col_count == 3:
        total = sum(weights)
        ratios = [weight / total for weight in weights]
        ratios[0] = min(max(ratios[0], 0.18), 0.28)
        ratios[1] = min(max(ratios[1], 0.38), 0.56)
        ratios[2] = max(0.20, 1.0 - ratios[0] - ratios[1])
        ratio_total = sum(ratios)
        weights = [ratio / ratio_total for ratio in ratios]
    else:
        total = sum(weights)
        weights = [weight / total for weight in weights]

    widths = []
    for weight in weights:
        widths.append(max(0.70 * inch, available_width * weight))
    scale = available_width / max(sum(widths), 1)
    return [width * scale for width in widths]


def build_pdf_table(rows: Sequence[Sequence[str]], styles: dict, available_width: float, *, long_table: bool):
    widths = approximate_table_widths(rows, available_width)
    formatted_rows = []
    for row_idx, row in enumerate(rows):
        formatted_row = []
        style = styles["table_header"] if row_idx == 0 else styles["table_cell"]
        for value in row:
            formatted_row.append(Paragraph(to_reportlab_markup(value), style))
        formatted_rows.append(formatted_row)

    table_cls = LongTable if long_table else Table
    table = table_cls(formatted_rows, colWidths=widths, repeatRows=1)
    table.hAlign = "LEFT"
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FBFD")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def build_pdf_image(source: Path, block: dict, available_width: float, styles: dict):
    image_path = resolve_asset_path(source, block["path"])
    if not image_path.exists():
        return Paragraph(f"<i>Missing figure asset:</i> {escape(block['path'])}", styles["body"])

    width, height = ImageReader(str(image_path)).getSize()
    max_width = available_width
    max_height = 6.8 * inch
    scale = min(max_width / max(width, 1), max_height / max(height, 1), 1.0)
    image = RLImage(str(image_path), width=width * scale, height=height * scale)
    image.hAlign = "CENTER"
    caption = Paragraph(to_reportlab_markup(block["alt"]), styles["caption"])
    return KeepTogether([image, Spacer(1, 0.06 * inch), caption])


def apply_pdf_metadata(canvas, doc, title: str) -> None:
    canvas.setTitle(title)
    canvas.setAuthor(DOCUMENT_AUTHOR)
    canvas.setSubject(document_subject(Path(doc.filename).with_suffix(".md")))
    canvas.setCreator("Footware-Design artifact renderer")
    canvas.setKeywords("footwear, LCA, sustainability, manuscript")


def render_markdown_to_pdf_reportlab(source: Path, output: Path) -> None:
    styles = build_pdf_styles()
    doc = SimpleDocTemplate(
        str(output),
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    flowables = []
    available_width = LETTER[0] - doc.leftMargin - doc.rightMargin

    for block in parse_markdown_blocks(source.read_text(encoding="utf-8")):
        btype = block["type"]
        if btype == "heading":
            style = styles[f"heading{block['level']}"]
            flowables.append(Paragraph(to_reportlab_markup(block["text"]), style))
        elif btype == "paragraph":
            flowables.append(Paragraph(to_reportlab_markup(block["text"]), styles["body"]))
        elif btype == "list":
            for index, item in enumerate(block["items"], start=1):
                bullet_text = f"{index}." if block["ordered"] else "&bull;"
                flowables.append(
                    Paragraph(f"{bullet_text} {to_reportlab_markup(item)}", styles["bullet"])
                )
        elif btype == "code":
            flowables.append(Preformatted(block["text"], styles["code"]))
            flowables.append(Spacer(1, 0.08 * inch))
        elif btype == "table":
            rows = [block["header"]] + block["rows"]
            if len(rows) <= 8:
                table = build_pdf_table(rows, styles, available_width, long_table=False)
                flowables.append(KeepTogether([table, Spacer(1, 0.12 * inch)]))
            else:
                table = build_pdf_table(rows, styles, available_width, long_table=True)
                flowables.append(table)
                flowables.append(Spacer(1, 0.12 * inch))
        elif btype == "image":
            flowables.append(build_pdf_image(source, block, available_width, styles))
            flowables.append(Spacer(1, 0.08 * inch))
        elif btype == "rule":
            flowables.append(Spacer(1, 0.08 * inch))

    output.parent.mkdir(parents=True, exist_ok=True)
    title = document_title(source)
    doc.build(
        flowables,
        onFirstPage=lambda canvas, built_doc: apply_pdf_metadata(canvas, built_doc, title),
        onLaterPages=lambda canvas, built_doc: apply_pdf_metadata(canvas, built_doc, title),
    )


def render_markdown_to_pdf_chrome(source: Path, output: Path) -> None:
    chrome_binary = find_chrome_binary()
    if chrome_binary is None:
        raise RuntimeError("Google Chrome is not available for tagged PDF export.")

    title = document_title(source)
    blocks = parse_markdown_blocks(source.read_text(encoding="utf-8"))
    html = markdown_blocks_to_html(blocks, title, source)

    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile("w", suffix=".html", delete=False, dir=output.parent, encoding="utf-8") as handle:
        temp_html = Path(handle.name)
        handle.write(html)

    try:
        file_url = temp_html.resolve().as_uri()
        subprocess.run(
            [
                chrome_binary,
                "--headless=new",
                "--no-sandbox",
                "--disable-gpu",
                "--no-pdf-header-footer",
                f"--print-to-pdf={output}",
                file_url,
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    finally:
        temp_html.unlink(missing_ok=True)


def render_markdown_to_pdf(source: Path, output: Path) -> None:
    try:
        render_markdown_to_pdf_chrome(source, output)
    except Exception:
        render_markdown_to_pdf_reportlab(source, output)


def csv_rows(path: Path) -> List[List[str]]:
    with path.open(newline="", encoding="utf-8") as handle:
        return [row for row in csv.reader(handle)]


def autosize_worksheet(worksheet) -> None:
    for col_idx in range(1, worksheet.max_column + 1):
        max_length = 0
        for cell in worksheet[get_column_letter(col_idx)]:
            value = "" if cell.value is None else str(cell.value)
            max_length = max(max_length, len(value))
        worksheet.column_dimensions[get_column_letter(col_idx)].width = min(max(max_length + 2, 10), 40)


def write_rows_to_sheet(workbook: Workbook, title: str, rows: Iterable[Sequence[str]]) -> None:
    worksheet = workbook.create_sheet(title)
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    header_font = Font(bold=True)
    wrap = Alignment(vertical="top", wrap_text=True)

    for row_idx, row in enumerate(rows, start=1):
        worksheet.append(list(row))
        for cell in worksheet[row_idx]:
            cell.alignment = wrap
            if row_idx == 1:
                cell.fill = header_fill
                cell.font = header_font
    worksheet.freeze_panes = "A2"
    worksheet.auto_filter.ref = worksheet.dimensions
    autosize_worksheet(worksheet)


def render_csv_to_xlsx(source: Path, output: Path) -> None:
    workbook = Workbook()
    default_sheet = workbook.active
    workbook.remove(default_sheet)
    write_rows_to_sheet(workbook, source.stem[:31], csv_rows(source))
    output.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output)


def render_shoe_full_lca(output: Path) -> None:
    source = ROOT / "statistical_analysis_results.csv"
    with source.open(newline="", encoding="utf-8") as handle:
        rows = list(csv.DictReader(handle))

    workbook = Workbook()
    default_sheet = workbook.active
    workbook.remove(default_sheet)

    per_shoe_rows = [
        [
            "Model",
            "Materials",
            "Material_kgCO2e_per_pair",
            "Manufacturing_kgCO2e_per_pair",
            "Transport_kgCO2e_per_pair",
            "Packaging_kgCO2e_per_pair",
            "EndOfLife_kgCO2e_per_pair",
            "Total_LCA_kgCO2e_per_pair",
        ]
    ]
    for row in rows:
        material = float(row["Material_kgCO2e_per_pair"])
        per_shoe_rows.append(
            [
                row["Model"],
                row["Materials"],
                f"{material:.5f}",
                f"{material * 0.25:.5f}",
                f"{material * 0.15:.5f}",
                f"{material * 0.05:.5f}",
                f"{material * 0.02:.5f}",
                f"{float(row['Total_LCA_kgCO2e_per_pair']):.5f}",
            ]
        )
    write_rows_to_sheet(workbook, "Per-shoe LCA", per_shoe_rows)

    assumptions_rows = [
        ["Component", "Share_vs_Materials", "Notes"],
        ["Materials", "1.00", "Calculated from emission factor and estimated component masses."],
        ["Manufacturing", "0.25", "Assumed 25% of material footprint."],
        ["Transport", "0.15", "Assumed 15% of material footprint."],
        ["Packaging", "0.05", "Assumed 5% of material footprint."],
        ["End-of-life", "0.02", "Assumed 2% of material footprint."],
    ]
    write_rows_to_sheet(workbook, "Assumptions", assumptions_rows)

    output.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output)


def main() -> None:
    analyze_expanded_dataset()
    generate_figure_set()

    for item in MARKDOWN_ARTIFACTS:
        render_markdown_to_docx(item["source"], item["docx"])
        if item["pdf"] is not None:
            render_markdown_to_pdf(item["source"], item["pdf"])

    for source, output in CSV_XLSX_ARTIFACTS:
        render_csv_to_xlsx(source, output)

    render_shoe_full_lca(ROOT / "shoe_full_LCA.xlsx")


if __name__ == "__main__":
    main()
